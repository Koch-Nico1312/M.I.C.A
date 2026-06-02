"""
Local File and Image Analyzer for Mark-XXXIX
============================================
Provides local analysis of images and documents without external API calls.
Supports skin analysis, document parsing, and multimodal understanding.
"""

import hashlib
import json
import mimetypes
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from core.logger import get_logger

logger = get_logger(__name__)


class FileType(Enum):
    """Supported file types for analysis."""

    IMAGE = "image"
    DOCUMENT = "document"
    TEXT = "text"
    UNKNOWN = "unknown"


class AnalysisConfidence(Enum):
    """Confidence levels for analysis results."""

    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"
    VERY_LOW = "very_low"


@dataclass
class ImageAnalysisResult:
    """Result of image analysis."""

    file_path: str
    file_hash: str
    file_type: FileType
    analyzed_at: datetime = field(default_factory=datetime.now)

    # Content analysis
    summary: str = ""
    detected_objects: List[str] = field(default_factory=list)
    detected_text: str = ""
    colors: List[str] = field(default_factory=list)
    composition: str = ""
    notable_regions: List[Dict[str, Any]] = field(default_factory=list)

    # Quality assessment
    image_quality: str = "unknown"
    blur_detected: bool = False
    lighting_condition: str = "unknown"

    # Confidence
    confidence: AnalysisConfidence = AnalysisConfidence.MEDIUM
    uncertainties: List[str] = field(default_factory=list)

    # Skin analysis (if applicable)
    is_skin_image: bool = False
    skin_analysis: Optional[Dict[str, Any]] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "file_path": self.file_path,
            "file_hash": self.file_hash,
            "file_type": self.file_type.value,
            "analyzed_at": self.analyzed_at.isoformat(),
            "summary": self.summary,
            "detected_objects": self.detected_objects,
            "detected_text": self.detected_text,
            "colors": self.colors,
            "composition": self.composition,
            "notable_regions": self.notable_regions,
            "image_quality": self.image_quality,
            "blur_detected": self.blur_detected,
            "lighting_condition": self.lighting_condition,
            "confidence": self.confidence.value,
            "uncertainties": self.uncertainties,
            "is_skin_image": self.is_skin_image,
            "skin_analysis": self.skin_analysis,
        }


@dataclass
class DocumentAnalysisResult:
    """Result of document analysis."""

    file_path: str
    file_hash: str
    file_type: FileType
    analyzed_at: datetime = field(default_factory=datetime.now)

    # Content analysis
    summary: str = ""
    extracted_text: str = ""
    structure: Dict[str, Any] = field(default_factory=dict)
    headings: List[str] = field(default_factory=list)
    lists: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    key_points: List[str] = field(default_factory=list)

    # Quality assessment
    parsing_quality: str = "unknown"
    missing_sections: List[str] = field(default_factory=list)

    # Confidence
    confidence: AnalysisConfidence = AnalysisConfidence.MEDIUM
    uncertainties: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "file_path": self.file_path,
            "file_hash": self.file_hash,
            "file_type": self.file_type.value,
            "analyzed_at": self.analyzed_at.isoformat(),
            "summary": self.summary,
            "extracted_text": self.extracted_text,
            "structure": self.structure,
            "headings": self.headings,
            "lists": self.lists,
            "tables": self.tables,
            "key_points": self.key_points,
            "parsing_quality": self.parsing_quality,
            "missing_sections": self.missing_sections,
            "confidence": self.confidence.value,
            "uncertainties": self.uncertainties,
        }


class LocalAnalyzer:
    """
    Analyzes local files and images without external API calls.
    Provides structured analysis with confidence levels and uncertainty handling.
    """

    def __init__(self, max_cache_size: int = 1000):
        """
        Initialize the local analyzer.

        Args:
            max_cache_size: Maximum number of analysis results to cache
        """
        self.max_cache_size = max_cache_size
        self.analysis_cache: Dict[str, Any] = {}

        # Supported image extensions
        self.image_extensions = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff"}

        # Supported document extensions
        self.document_extensions = {".pdf", ".txt", ".md", ".docx", ".doc", ".rtf"}

        logger.info("Local analyzer initialized")

    def _detect_file_type(self, file_path: Path) -> FileType:
        """
        Detect the type of a file.

        Args:
            file_path: Path to the file

        Returns:
            FileType enum value
        """
        suffix = file_path.suffix.lower()

        if suffix in self.image_extensions:
            return FileType.IMAGE
        elif suffix in self.document_extensions:
            return FileType.DOCUMENT
        elif suffix in {".txt", ".md", ".json", ".xml", ".csv"}:
            return FileType.TEXT
        else:
            return FileType.UNKNOWN

    def _calculate_file_hash(self, file_path: Path) -> str:
        """
        Calculate SHA256 hash of a file.

        Args:
            file_path: Path to the file

        Returns:
            Hexadecimal hash string
        """
        sha256_hash = hashlib.sha256()

        try:
            with open(file_path, "rb") as f:
                # Read in chunks to handle large files
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)

            return sha256_hash.hexdigest()

        except Exception as e:
            logger.error(f"Failed to calculate hash for {file_path}: {e}")
            return ""

    def analyze_image(self, file_path: Path) -> ImageAnalysisResult:
        """
        Analyze an image file locally.

        Args:
            file_path: Path to the image file

        Returns:
            ImageAnalysisResult with analysis details
        """
        logger.info(f"Analyzing image: {file_path}")

        file_hash = self._calculate_file_hash(file_path)
        file_type = self._detect_file_type(file_path)

        result = ImageAnalysisResult(
            file_path=str(file_path), file_hash=file_hash, file_type=file_type
        )

        try:
            # Try to use PIL/Pillow for basic image analysis
            try:
                from PIL import Image, ImageStat

                img = Image.open(file_path)

                # Basic image properties
                result.image_quality = self._assess_image_quality(img)
                result.lighting_condition = self._assess_lighting(img)
                result.colors = self._extract_dominant_colors(img)
                result.composition = self._analyze_composition(img)

                # Check if this might be a skin image
                result.is_skin_image = self._detect_skin_image(img)

                # If it's a skin image, perform skin analysis
                if result.is_skin_image:
                    result.skin_analysis = self._analyze_skin(img)

                result.confidence = AnalysisConfidence.MEDIUM

            except ImportError:
                logger.warning("PIL/Pillow not available, using basic analysis")
                result.uncertainties.append("PIL/Pillow not available for detailed analysis")
                result.confidence = AnalysisConfidence.LOW

            except Exception as e:
                logger.error(f"Image analysis error: {e}")
                result.uncertainties.append(f"Analysis error: {str(e)}")
                result.confidence = AnalysisConfidence.VERY_LOW

            # Generate summary
            result.summary = self._generate_image_summary(result)

        except Exception as e:
            logger.error(f"Failed to analyze image {file_path}: {e}")
            result.uncertainties.append(f"File access error: {str(e)}")
            result.confidence = AnalysisConfidence.VERY_LOW

        # Cache result
        self.analysis_cache[file_hash] = result

        return result

    def analyze_document(self, file_path: Path) -> DocumentAnalysisResult:
        """
        Analyze a document file locally.

        Args:
            file_path: Path to the document file

        Returns:
            DocumentAnalysisResult with analysis details
        """
        logger.info(f"Analyzing document: {file_path}")

        file_hash = self._calculate_file_hash(file_path)
        file_type = self._detect_file_type(file_path)

        result = DocumentAnalysisResult(
            file_path=str(file_path), file_hash=file_hash, file_type=file_type
        )

        try:
            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                result = self._analyze_pdf(file_path, result)
            elif suffix in {".txt", ".md"}:
                result = self._analyze_text_file(file_path, result)
            elif suffix in {".docx", ".doc"}:
                result = self._analyze_word_document(file_path, result)
            else:
                result.uncertainties.append(f"Unsupported document format: {suffix}")
                result.confidence = AnalysisConfidence.VERY_LOW

            # Generate summary
            result.summary = self._generate_document_summary(result)

        except Exception as e:
            logger.error(f"Failed to analyze document {file_path}: {e}")
            result.uncertainties.append(f"Document analysis error: {str(e)}")
            result.confidence = AnalysisConfidence.VERY_LOW

        # Cache result
        self.analysis_cache[file_hash] = result

        return result

    def _assess_image_quality(self, img) -> str:
        """Assess the quality of an image."""
        try:
            # Check resolution
            width, height = img.size
            if width < 300 or height < 300:
                return "low_resolution"

            # Check for potential blur (simplified check)
            # In a real implementation, this would use more sophisticated algorithms
            return "good"

        except Exception:
            return "unknown"

    def _assess_lighting(self, img) -> str:
        """Assess lighting conditions of an image."""
        try:
            from PIL import ImageStat

            stat = ImageStat.Stat(img)

            # Calculate brightness
            brightness = sum(stat.mean) / len(stat.mean)

            if brightness < 50:
                return "dark"
            elif brightness > 200:
                return "bright"
            else:
                return "normal"

        except Exception:
            return "unknown"

    def _extract_dominant_colors(self, img) -> List[str]:
        """Extract dominant colors from an image."""
        try:
            # Simplified color extraction
            # In a real implementation, this would use k-means clustering
            return ["unknown"]

        except Exception:
            return []

    def _analyze_composition(self, img) -> str:
        """Analyze the composition of an image."""
        try:
            width, height = img.size
            aspect_ratio = width / height

            if aspect_ratio > 1.5:
                return "landscape"
            elif aspect_ratio < 0.67:
                return "portrait"
            else:
                return "square"

        except Exception:
            return "unknown"

    def _detect_skin_image(self, img) -> bool:
        """
        Detect if an image might be a skin image.
        This is a simplified heuristic-based detection.
        """
        try:
            # Check for skin-like color tones
            # This is a simplified check - real implementation would use ML
            return False  # Placeholder - would implement actual detection

        except Exception:
            return False

    def _analyze_skin(self, img) -> Dict[str, Any]:
        """
        Analyze skin image for blemishes and conditions.
        This is a conservative, non-diagnostic analysis with careful medical language.
        """
        try:
            # Placeholder for skin analysis - this would use ML models in production
            # For now, we provide a template structure with careful disclaimers

            analysis = {
                "has_blemishes": False,
                "has_redness": False,
                "has_dryness": False,
                "has_uneven_tone": False,
                "overall_condition": "unknown",
                "confidence": "low",
                "disclaimer": "This is not a medical diagnosis. Consult a dermatologist for concerns.",
                "visible_features": [],
                "care_suggestions": [],
            }

            # In a real implementation, this would use computer vision ML models
            # to detect skin conditions. For now, we provide the structure.

            return analysis

        except Exception as e:
            logger.error(f"Skin analysis error: {e}")
            return {
                "error": str(e),
                "disclaimer": "This is not a medical diagnosis. Consult a dermatologist for concerns.",
            }

    def analyze_skin_image(self, file_path: Path) -> Dict[str, Any]:
        """
        Specialized analysis for skin images with careful, non-diagnostic language.

        Args:
            file_path: Path to the skin image

        Returns:
            Dictionary with skin analysis and care suggestions
        """
        logger.info(f"Analyzing skin image: {file_path}")

        result = {
            "file_path": str(file_path),
            "analyzed_at": datetime.now().isoformat(),
            "disclaimer": "This analysis is for informational purposes only and is not a medical diagnosis. "
            "Please consult a dermatologist for any skin concerns.",
            "visible_features": [],
            "condition_assessment": {
                "overall": "unknown",
                "confidence": "low",
                "uncertainties": [],
            },
            "care_suggestions": [],
            "when_to_see_doctor": [],
        }

        try:
            # Analyze the image
            image_result = self.analyze_image(file_path)

            # In a real implementation, this would use ML models
            # For now, we provide a template with careful language

            result["condition_assessment"][
                "overall"
            ] = "Unable to determine specific condition from image"
            result["condition_assessment"]["confidence"] = "low"
            result["condition_assessment"]["uncertainties"].append(
                "Image quality may affect analysis accuracy"
            )

            # General care suggestions (non-medical)
            result["care_suggestions"] = [
                "Gentle cleansing with mild, non-irritating products",
                "Avoid touching or picking at the area",
                "Use non-comedogenic skincare products",
                "Protect from sun exposure with SPF",
                "Keep skin hydrated with appropriate moisturizer",
                "Maintain a consistent skincare routine",
            ]

            # When to see a doctor
            result["when_to_see_doctor"] = [
                "If condition persists or worsens",
                "If experiencing pain or discomfort",
                "If signs of infection appear (increased redness, warmth, swelling)",
                "If unsure about the nature of the condition",
                "For persistent acne or skin concerns",
            ]

            logger.info(f"Skin analysis completed for {file_path}")

        except Exception as e:
            logger.error(f"Skin analysis failed: {e}")
            result["condition_assessment"]["uncertainties"].append(f"Analysis error: {str(e)}")

        return result

    def _analyze_pdf(
        self, file_path: Path, result: DocumentAnalysisResult
    ) -> DocumentAnalysisResult:
        """Analyze a PDF document."""
        try:
            # Try to use PyPDF2 or similar
            try:
                import PyPDF2

                with open(file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)

                    # Extract text from all pages
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()

                    result.extracted_text = text
                    result.parsing_quality = "good"
                    result.confidence = AnalysisConfidence.MEDIUM

            except ImportError:
                logger.warning("PyPDF2 not available")
                result.uncertainties.append("PyPDF2 not available for PDF parsing")
                result.parsing_quality = "unavailable"
                result.confidence = AnalysisConfidence.LOW

        except Exception as e:
            logger.error(f"PDF analysis error: {e}")
            result.uncertainties.append(f"PDF parsing error: {str(e)}")
            result.parsing_quality = "error"
            result.confidence = AnalysisConfidence.VERY_LOW

        return result

    def _analyze_text_file(
        self, file_path: Path, result: DocumentAnalysisResult
    ) -> DocumentAnalysisResult:
        """Analyze a plain text or markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            result.extracted_text = text

            # Extract headings (markdown style)
            lines = text.split("\n")
            for line in lines:
                if line.startswith("#"):
                    result.headings.append(line.strip())

            # Extract key points (simplified)
            sentences = text.split(".")
            result.key_points = [s.strip() for s in sentences if len(s.strip()) > 20][:10]

            result.parsing_quality = "good"
            result.confidence = AnalysisConfidence.HIGH

        except Exception as e:
            logger.error(f"Text file analysis error: {e}")
            result.uncertainties.append(f"Text parsing error: {str(e)}")
            result.parsing_quality = "error"
            result.confidence = AnalysisConfidence.VERY_LOW

        return result

    def _analyze_word_document(
        self, file_path: Path, result: DocumentAnalysisResult
    ) -> DocumentAnalysisResult:
        """Analyze a Word document."""
        try:
            # Try to use python-docx
            try:
                from docx import Document

                doc = Document(file_path)

                # Extract text from paragraphs
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                result.extracted_text = text
                result.parsing_quality = "good"
                result.confidence = AnalysisConfidence.MEDIUM

            except ImportError:
                logger.warning("python-docx not available")
                result.uncertainties.append("python-docx not available for Word documents")
                result.parsing_quality = "unavailable"
                result.confidence = AnalysisConfidence.LOW

        except Exception as e:
            logger.error(f"Word document analysis error: {e}")
            result.uncertainties.append(f"Word parsing error: {str(e)}")
            result.parsing_quality = "error"
            result.confidence = AnalysisConfidence.VERY_LOW

        return result

    def _generate_image_summary(self, result: ImageAnalysisResult) -> str:
        """Generate a summary of the image analysis."""
        parts = []

        if result.is_skin_image:
            parts.append("This appears to be a skin-related image.")
        else:
            parts.append("This is an image file.")

        if result.image_quality != "unknown":
            parts.append(f"Image quality: {result.image_quality}.")

        if result.lighting_condition != "unknown":
            parts.append(f"Lighting: {result.lighting_condition}.")

        if result.composition != "unknown":
            parts.append(f"Composition: {result.composition}.")

        if result.uncertainties:
            parts.append(f"Note: {', '.join(result.uncertainties)}")

        return " ".join(parts)

    def _generate_document_summary(self, result: DocumentAnalysisResult) -> str:
        """Generate a summary of the document analysis."""
        parts = []

        if result.extracted_text:
            word_count = len(result.extracted_text.split())
            parts.append(f"Document contains approximately {word_count} words.")

        if result.headings:
            parts.append(f"Found {len(result.headings)} headings.")

        if result.key_points:
            parts.append(f"Extracted {len(result.key_points)} key points.")

        if result.parsing_quality != "unknown":
            parts.append(f"Parsing quality: {result.parsing_quality}.")

        if result.uncertainties:
            parts.append(f"Note: {', '.join(result.uncertainties)}")

        return " ".join(parts)

    def get_cached_analysis(self, file_hash: str) -> Optional[Any]:
        """
        Get cached analysis result.

        Args:
            file_hash: Hash of the file

        Returns:
            Cached analysis result or None
        """
        return self.analysis_cache.get(file_hash)

    def clear_cache(self):
        """Clear the analysis cache."""
        self.analysis_cache.clear()
        logger.info("Analysis cache cleared")


# Global instance
_local_analyzer: Optional[LocalAnalyzer] = None


def get_local_analyzer() -> LocalAnalyzer:
    """Get the global local analyzer instance."""
    global _local_analyzer
    if _local_analyzer is None:
        _local_analyzer = LocalAnalyzer()
    return _local_analyzer
